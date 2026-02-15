"""
✨ Project: Markdown to Docx Converter
✨ Description: A batch converter for Markdown/Text files to Word documents.
✨ Coded by Ajin (Gemini) with ❤️ 
"""

import os
import re
import sys

try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    print("❌ 缺少工具包！请先运行：pip install python-docx")
    sys.exit()

# --- 配置区 (可根据喜好调整) ---
SIGNATURE = "Coded by Ajin (Gemini) with ❤️"
BOLD_COLOR = RGBColor(0, 51, 102)  # 深蓝色
EXCLUDED_CHARS_PATTERN = r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]'

def clean_xml_string(content):
    """
    清洗字符串，剔除不符合 XML 规范的控制字符，防止 ValueError
    """
    if not isinstance(content, str):
        return str(content)
    return re.sub(EXCLUDED_CHARS_PATTERN, '', content)

def markdown_to_docx(file_path):
    base_name = os.path.splitext(file_path)[0]
    output_path = f"{base_name}.docx"
    
    print(f"📄 正在精洗并转换: {os.path.basename(file_path)} ...")

    doc = Document()
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print(f"⚠️ 读取失败 {file_path}: {e}")
        return

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("") 
            continue

        # 处理标题
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            content = line.lstrip('#').strip()
            level = min(level, 9)
            doc.add_heading(clean_xml_string(content), level=level)
        
        # 处理普通段落及加粗逻辑
        else:
            p = doc.add_paragraph()
            # 这里的正则保留了 ** 标记以便后续识别
            parts = re.split(r'(\*\*.*?\*\*)', line)
            
            for part in parts:
                clean_part = clean_xml_string(part)
                if part.startswith('**') and part.endswith('**'):
                    text = clean_part[2:-2]
                    run = p.add_run(text)
                    run.bold = True
                    run.font.color.rgb = BOLD_COLOR
                else:
                    p.add_run(clean_part)

    try:
        doc.save(output_path)
        print(f"✅ 成功生成: {os.path.basename(output_path)}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")

def process_path(path):
    path = path.strip().strip('"').strip("'")
    
    if not os.path.exists(path):
        print(f"❓ 找不到路径: {path}")
        return

    if os.path.isfile(path):
        if path.lower().endswith(('.md', '.txt')):
            markdown_to_docx(path)
    elif os.path.isdir(path):
        print(f"📁 正在扫描文件夹: {path}")
        for root, _, files in os.walk(path):
            for file in files:
                if file.lower().endswith(('.md', '.txt')):
                    full_path = os.path.join(root, file)
                    markdown_to_docx(full_path)

def main():
    print("="*50)
    print(f" 🚀 {SIGNATURE}")
    print(" ✨ Markdown 转 Word 批量清洗车间 ✨")
    print("="*50)
    print("\n💡 使用方法：")
    print("1. 直接把文件或文件夹「拖入」这个窗口")
    print("2. 按回车键确认开始转换 (输入 q 退出)\n")

    while True:
        user_input = input("👉 请拖入路径: ").strip()
        
        if user_input.lower() == 'q':
            print("👋 下班啦，期待下次为你服务！")
            break
            
        if not user_input:
            continue

        # 处理 Windows 多路径拖入
        if '"' in user_input:
            paths = re.findall(r'"(.*?)"', user_input)
            if not paths:
                paths = [user_input]
        else:
            paths = [user_input]

        for p in paths:
            process_path(p)
        
        print("\n✨ 这一波处理完啦！可以继续拖入。")
        print("-" * 30)

if __name__ == "__main__":
    main()